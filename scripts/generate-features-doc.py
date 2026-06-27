from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ORANGE  = RGBColor(0xCC, 0x55, 0x00)
INK     = RGBColor(0x0A, 0x0A, 0x0A)
GREY    = RGBColor(0x66, 0x66, 0x66)
RULE    = RGBColor(0xE0, 0xE0, 0xDA)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BGLIGHT = RGBColor(0xFD, 0xF6, 0xF0)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, color=INK, size=11, font='Courier New'):
    run = para.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = font
    return run

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph spacing ─────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Courier New'
style.font.size = Pt(10)
style.font.color.rgb = INK
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after  = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER / HEADER
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('rekōdo')
run.font.name   = 'Georgia'
run.font.size   = Pt(36)
run.font.bold   = True
run.font.color.rgb = ORANGE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('rekodo.co  ·  Product Features')
run2.font.name  = 'Courier New'
run2.font.size  = Pt(10)
run2.font.color.rgb = GREY

# Horizontal rule (via a table with a bottom border)
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
cell = tbl.rows[0].cells[0]
cell.text = ''
set_cell_border(cell, bottom={'val': 'single', 'sz': '8', 'color': 'CC5500'})
set_cell_bg(cell, 'FFFFFF')
cell.paragraphs[0].paragraph_format.space_before = Pt(0)
cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.name      = 'Courier New'
    run.font.size      = Pt(9)
    run.font.bold      = True
    run.font.color.rgb = ORANGE
    run.font.all_caps  = False
    # underline via character spacing trick — just use color + caps
    return p

def sub_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name      = 'Georgia'
    run.font.size      = Pt(12)
    run.font.bold      = True
    run.font.color.rgb = INK
    return p

def supporter_badge(para):
    run = para.add_run('  SUPPORTER')
    run.font.name      = 'Courier New'
    run.font.size      = Pt(7)
    run.font.bold      = True
    run.font.color.rgb = ORANGE

def bullet(text, indent=0, supporter=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    run = p.add_run(f'— {text}')
    run.font.name      = 'Courier New'
    run.font.size      = Pt(10)
    run.font.color.rgb = INK
    if supporter:
        supporter_badge(p)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.name      = 'Courier New'
    run.font.size      = Pt(8.5)
    run.font.italic    = True
    run.font.color.rgb = GREY
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ── COLLECTION ────────────────────────────────────────────────────────────────
section_heading('Collection')

sub_heading('Import')
bullet('Discogs OAuth sync — read-only, never writes back to Discogs')
bullet('Discogs CSV export upload — Backup / Bulk Import fallback when API is unavailable')
bullet('Bandcamp digital collection import via username (linked from your profile)')

sub_heading('Browsing & Search')
bullet('A–Z record list with album art thumbnails')
bullet('Search by artist, album, or label')
bullet('Filter by Genre, Decade, Format, Desirability, or Feeling')
bullet('Sort by name, value, or year')

sub_heading('Record Detail Panel')
bullet('Label, Format, Edition size (where available), Vinyl colour')
bullet('Media & sleeve condition grades')
bullet('Country of pressing, Year, Genre, Style, Cat #, Producers')
bullet('Discogs marketplace value — lowest listing and last sold price')
bullet('Desirability tier — Rare / Cult Pressing / In Demand / Widely Loved')
bullet('Open to Offers flag — surfaces on public profile and Sell List')
bullet('Essential tag')
bullet('Feeling tag — how does this record make you feel')
bullet('Memory notes — personal, optionally shareable')
bullet('Tracklist')
bullet('Bandcamp track lookup')

sub_heading('Collection Value')
bullet('Total low / median / high valuation from Discogs marketplace data')

# ── INSIGHTS ──────────────────────────────────────────────────────────────────
section_heading('Insights')

sub_heading('Collection Tab  —  Free')
bullet('Daily Pick')
bullet('On This Day')
bullet('Essentials Wall')
bullet('Feeling breakdown')
bullet('Genre, Style, Country, Format, Vinyl colour breakdowns')
bullet('Top Artists, Labels, Producers')
bullet('Desirability breakdown')
bullet('Top records by value')
bullet('Media & sleeve condition breakdown')
bullet('Top played records & play style breakdown')

sub_heading('Taste Profile Tab')
bullet('7 Spectrum Dimensions — Canon ↔ Obscure, Nostalgic ↔ Contemporary,', supporter=True)
bullet('Completist, Vinyl Pure ↔ Format Agnostic, Ambient ↔ Abrasive, and more', indent=1)
bullet('Each dimension derived from real collection data, not a quiz', indent=1)

# ── DIG ───────────────────────────────────────────────────────────────────────
section_heading('Dig  —  AI-Powered Discovery')

bullet('Discover — personalised record recommendations for artists you don\'t own yet, based on your collection, genres, labels, and lists')
bullet('Explore — surfaces hidden gems from within your own collection that deserve more attention')
bullet('Style Dig — recommendations filtered to a specific style or scene')
bullet('Dig History — past sessions saved locally')
bullet('Add recommendations directly to Wantlist')
bullet('Quick links to Apple Music, Spotify, Tidal, Discogs, Bandcamp, Rough Trade, Juno, Boomkat, eBay')
note('Free: 3 Dig sessions per day  ·  Supporter: unlimited')

# ── DEEP DIVE ─────────────────────────────────────────────────────────────────
section_heading('Deep Dive')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('Available to rekōdo Supporters')
run.font.name      = 'Courier New'
run.font.size      = Pt(9)
run.font.italic    = True
run.font.color.rgb = ORANGE

bullet('Full artist deep dive for any artist in your collection')
bullet('Every record you own by that artist with pressing details and market values')
bullet('Artist image and biographical context')
bullet('Discography gaps and related releases')

sub_heading('Deep Dive Tabs')
bullet('Essential Albums — AI-ranked must-have records for the artist')
bullet('Podcasts — artist-specific episodes surfaced from Apple Podcasts & Spotify')
bullet('Books & Audiobooks — books about the artist linked to Amazon / Audible')
bullet('Interviews — key interviews and profiles')
bullet('Related Artists — connected artists based on genre, label, and era')
bullet('Blind Spot — records by this artist you don\'t own yet')

# ── SELECTS ───────────────────────────────────────────────────────────────────
section_heading('Selects')

bullet('Artist Spotlight — rotating editorial feature on a key artist and their catalogue')
bullet('Label Spotlight — deep dive into a label\'s history, sound, and landmark releases')
bullet('New Releases — releases, represses, and preorders from labels you subscribe to via email')
bullet('Live — upcoming gigs near you featuring artists from your collection (Ticketmaster)')

# ── ARCHETYPES ────────────────────────────────────────────────────────────────
section_heading('Archetypes')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('Available to rekōdo Supporters')
run.font.name      = 'Courier New'
run.font.size      = Pt(9)
run.font.italic    = True
run.font.color.rgb = ORANGE

bullet('Collector archetype — calculated from your collection shape, genres, decades, labels, rarity distribution, and app usage')
bullet('Archetype essay — AI-generated written portrait of your collector identity')

# ── LISTS ─────────────────────────────────────────────────────────────────────
section_heading('Lists')

bullet('Create public or private lists')
bullet('Add records from your collection, search Discogs, or add individual songs')
bullet('Share public lists via a dedicated URL')
bullet('Collaborative discovery — browse other collectors\' public lists')

sub_heading('Playlist Generator')
bullet('Generate a playlist from your collection using a mood or prompt')
bullet('Tracks matched against your actual records and pulled into a Spotify playlist')
bullet('Save generated playlists and return to them later')
bullet('Share playlists with a public link')

# ── WANTLIST ──────────────────────────────────────────────────────────────────
section_heading('Wantlist  —  Private')

bullet('Auto-populated from Dig recommendations and Lists')
bullet('Discogs wantlist sync', supporter=True)
bullet('CSV import from Discogs', supporter=True)

# ── PROFILE ───────────────────────────────────────────────────────────────────
section_heading('Profile')

bullet('Public profile page with bio, city, star sign, and taste statement')
bullet('Top 5 records')
bullet('Essentials Wall — publicly viewable')
bullet('Public Lists')
bullet('Sell List — records flagged as Open to Offers')
bullet('Collection photos')
bullet('Connected accounts: Discogs, Spotify, Bandcamp')
bullet('Supporter golden ō badge (also available to donors)')

# ── FREE VS SUPPORTER TABLE ───────────────────────────────────────────────────
section_heading('Free vs Supporter')

doc.add_paragraph()

headers = ['Feature', 'Free', 'Supporter']
rows = [
    ('Collection sync & browsing',        '✓', '✓'),
    ('Insights — Collection tab',         '✓', '✓'),
    ('Dig',                               '3 / day', 'Unlimited'),
    ('Lists & Wantlist (from Dig/Lists)', '✓', '✓'),
    ('Selects & Live',                     '✓', '✓'),
    ('Insights — Taste Profile',          '—', '✓'),
    ('Deep Dive',                         '—', '✓'),
    ('Archetypes',                        '—', '✓'),
    ('Wantlist CSV import',               '—', '✓'),
    ('Discogs wantlist sync',             '—', '✓'),
    ('Golden ō badge',                    'Donor+', '✓'),
]

tbl = doc.add_table(rows=1 + len(rows), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = 'Table Grid'

# Header row
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, 'CC5500')
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name      = 'Courier New'
    run.font.size      = Pt(9)
    run.font.bold      = True
    run.font.color.rgb = WHITE

# Data rows
for ri, (feat, free, sup) in enumerate(rows):
    row   = tbl.rows[ri + 1]
    bg    = 'FDF6F0' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate([feat, free, sup]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = ORANGE if (ci > 0 and val == '✓') else INK

# Column widths
tbl.columns[0].width = Cm(9)
tbl.columns[1].width = Cm(3)
tbl.columns[2].width = Cm(3.5)

# ── FOOTER NOTE ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run('rekodo.co  ·  Independent. Ad-free. Built for serious collectors.')
run.font.name      = 'Courier New'
run.font.size      = Pt(8)
run.font.italic    = True
run.font.color.rgb = GREY

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = 'rekodo-features.docx'
doc.save(out)
print(f'Saved: {out}')
